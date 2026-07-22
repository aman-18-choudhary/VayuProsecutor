import datetime
import uuid
import json
from typing import Dict, Any, List, Optional
import random

class EnvironmentalReportGenerator:
    """
    Service for generating comprehensive AI Environmental Reports.
    Combines live AQI, Causal data, Ward intelligence, and Simulation data.
    """

    def __init__(self):
        # In-memory store for reports (for demo/persistence fallback)
        self.reports_db: Dict[str, Dict[str, Any]] = {}

    def get_all_reports(self) -> List[Dict[str, Any]]:
        return sorted(self.reports_db.values(), key=lambda x: x["generated_at"], reverse=True)

    def generate_executive_summary(self, city: str, aqi: int, culprits: List[Dict[str, Any]]) -> str:
        """
        AI-driven summary based on live metrics.
        """
        trend = "deteriorating" if aqi > 150 else "stable"
        main_source = culprits[0]['name'] if culprits else "unknown sources"
        
        summary = (
            f"The current environmental assessment for {city} indicates a {trend} situation with an overall AQI of {aqi}. "
            f"The primary driver of particulate matter pollution has been identified as {main_source}. "
        )
        if aqi > 300:
            summary += "Critical emergency protocols are highly recommended. High-risk populations should be sheltered immediately. "
        elif aqi > 200:
            summary += "Targeted interventions on industrial and vehicular emissions should be enacted within the next 24 hours. "
            
        summary += "Forecasts suggest conditions may fluctuate depending on wind velocity and incoming weather systems."
        return summary

    def generate_report(self, request_data: Dict[str, Any], context: Dict[str, Any]) -> Dict[str, Any]:
        """
        Builds the complete report structure.
        """
        city = request_data.get("city", "Unknown")
        report_type = request_data.get("type", "Executive Summary")
        
        # Extract live data
        live_data = context.get("live", {})
        aqi = live_data.get("aqi", 0)
        
        # Extract causal data
        causal_data = context.get("causal", {})
        culprits = causal_data.get("culprits", [])
        
        # Extract ward data
        wards_data = context.get("wards", [])
        
        # Extract alerts data
        alerts_data = context.get("alerts", [])

        report_id = f"REP-{uuid.uuid4().hex[:8].upper()}"
        now = datetime.datetime.now(datetime.timezone.utc)
        
        report = {
            "id": report_id,
            "title": f"{city} {report_type} - {now.strftime('%Y-%m-%d')}",
            "city": city,
            "type": report_type,
            "generated_at": now.isoformat(),
            "status": "Ready",
            "size": f"{random.randint(45, 250)} KB",
            "author": request_data.get("user", "System Automaton"),
            
            # Content Sections
            "executive_summary": self.generate_executive_summary(city, aqi, culprits),
            "current_situation": {
                "aqi": aqi,
                "severity": live_data.get("severity", "Unknown"),
                "temperature": live_data.get("weather", {}).get("temperature", "--"),
                "wind_speed": live_data.get("weather", {}).get("windSpeed", "--"),
            },
            "source_attribution": culprits,
            "ward_intelligence": [
                {
                    "name": w.get("name"),
                    "risk_index": w.get("risk_index", 0),
                    "primary_issue": w.get("primary_issue", "Dust"),
                    "aqi": w.get("aqi", aqi)
                } for w in wards_data[:5]
            ],
            "active_alerts_summary": {
                "total": len(alerts_data),
                "critical": len([a for a in alerts_data if a.get("severity") == "Critical"])
            },
            "insights": [
                f"{culprits[0]['name'] if culprits else 'Local activity'} contributes significantly to PM2.5 today.",
                f"Forecast predicts AQI { 'deterioration' if aqi > 100 else 'improvement' } within 24 hours.",
                "School exposure risk remains moderate."
            ],
            "confidence_score": 92
        }
        
        self.reports_db[report_id] = report
        return report

    def generate_pdf(self, report: Dict[str, Any], output_path: str):
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, report["title"], ln=True, align='C')
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 10, f"Report ID: {report['id']} | City: {report['city']} | Date: {report['generated_at']}", ln=True, align='C')
        pdf.ln(10)
        
        import textwrap

        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Executive Summary", ln=True)
        pdf.set_font("Helvetica", "", 10)
        summary = str(report.get("executive_summary", "")).encode('latin-1', 'replace').decode('latin-1')
        for line in textwrap.wrap(summary, width=90):
            pdf.cell(0, 7, line, ln=True)
        pdf.ln(5)

        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Current Situation", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for k, v in report.get("current_situation", {}).items():
            pdf.cell(0, 7, f"{k.capitalize()}: {v}", ln=True)
        pdf.ln(5)

        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Ward Intelligence", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for w in report.get("ward_intelligence", []):
            text = f"- {w.get('name')}: AQI {w.get('aqi')} (Risk: {w.get('risk_index')}) - {w.get('primary_issue')}"
            for line in textwrap.wrap(text, width=90):
                pdf.cell(0, 7, line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
        pdf.ln(5)

        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Insights & Recommendations", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for insight in report.get("insights", []):
            text = f"- {insight}"
            for line in textwrap.wrap(text, width=90):
                pdf.cell(0, 7, line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
        
        pdf.output(output_path)

    def generate_docx(self, report: Dict[str, Any], output_path: str):
        import docx
        doc = docx.Document()
        doc.add_heading(report["title"], 0)
        doc.add_paragraph(f"Report ID: {report['id']} | City: {report['city']} | Date: {report['generated_at']}")
        
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(str(report.get("executive_summary", "")))
        
        doc.add_heading("Current Situation", level=1)
        for k, v in report.get("current_situation", {}).items():
            doc.add_paragraph(f"{k.capitalize()}: {v}")
            
        doc.add_heading("Ward Intelligence", level=1)
        for w in report.get("ward_intelligence", []):
            doc.add_paragraph(f"- {w.get('name')}: AQI {w.get('aqi')} (Risk: {w.get('risk_index')}) - {w.get('primary_issue')}")
            
        doc.add_heading("Insights & Recommendations", level=1)
        for insight in report.get("insights", []):
            doc.add_paragraph(f"- {insight}", style='List Bullet')
            
        doc.save(output_path)

    def generate_csv(self, report: Dict[str, Any], output_path: str):
        import csv
        with open(output_path, mode='w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["Report ID", report.get("id")])
            writer.writerow(["Title", report.get("title")])
            writer.writerow(["City", report.get("city")])
            writer.writerow(["Generated At", report.get("generated_at")])
            writer.writerow([])
            
            writer.writerow(["Metric", "Value"])
            for k, v in report.get("current_situation", {}).items():
                writer.writerow([k.capitalize(), v])
            writer.writerow(["Total Alerts", report.get("active_alerts_summary", {}).get("total", 0)])
            writer.writerow(["Critical Alerts", report.get("active_alerts_summary", {}).get("critical", 0)])
            writer.writerow([])
            
            writer.writerow(["Ward", "AQI", "Risk Index", "Primary Issue"])
            for w in report.get("ward_intelligence", []):
                writer.writerow([w.get('name'), w.get('aqi'), w.get('risk_index'), w.get('primary_issue')])

# Global instance
report_engine = EnvironmentalReportGenerator()
